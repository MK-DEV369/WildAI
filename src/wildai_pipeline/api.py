from __future__ import annotations

import re
import sys
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[2]
_VENV_SITES = _ROOT / "venv" / "Lib" / "site-packages"
if _VENV_SITES.exists() and str(_VENV_SITES) not in sys.path:
    sys.path.insert(0, str(_VENV_SITES))

from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware

from .config import PipelineConfig
from .rag_engine import RAGEngine
from .schemas import BuildIndexResponse, QueryRequest, QueryResponse, SearchHit, SummaryRequest, ExportRequest
from .ollama_client import generate as ollama_generate
from fastapi.responses import FileResponse, JSONResponse
from tempfile import NamedTemporaryFile
import io
import datetime
from typing import Any
from fastapi.responses import StreamingResponse
import docx
import json
from pathlib import Path


HIGHLIGHT_STOPWORDS = {
    "a",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "for",
    "from",
    "how",
    "in",
    "is",
    "it",
    "of",
    "on",
    "or",
    "that",
    "the",
    "to",
    "was",
    "what",
    "when",
    "where",
    "which",
    "who",
    "why",
    "with",
}


def extract_highlight_terms(query: str, max_terms: int = 12) -> list[str]:
    """Return de-duplicated query tokens suitable for safe regex highlighting."""
    if not query.strip():
        return []

    # Keep words, numbers, and common internal separators (hyphen/apostrophe/slash).
    raw_terms = re.findall(r"[A-Za-z0-9]+(?:[-'/][A-Za-z0-9]+)*", query)

    unique_terms: list[str] = []
    seen_terms: set[str] = set()

    for term in raw_terms:
        normalized = term.lower()
        if len(normalized) < 2:
            continue
        if normalized in HIGHLIGHT_STOPWORDS:
            continue
        if normalized in seen_terms:
            continue

        seen_terms.add(normalized)
        unique_terms.append(term)

    unique_terms.sort(key=len, reverse=True)
    return unique_terms[:max_terms]


def create_app() -> FastAPI:
    config = PipelineConfig()
    engine = RAGEngine(config)

    app = FastAPI(title="WILDAI RAG API", version="0.1.0")
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

    @app.get("/api/health")
    def health() -> dict[str, str | bool]:
        return {"status": "ok", "index_ready": engine.index_path.exists()}

    @app.post("/api/index/rebuild", response_model=BuildIndexResponse)
    def rebuild_index() -> BuildIndexResponse:
        freq_path = config.output_dir / "word_frequencies.json"
        if freq_path.exists():
            try:
                freq_path.unlink()
            except Exception:
                pass
        payload = engine.build_index()
        return BuildIndexResponse(
            total_documents=int(payload["total_documents"]),
            total_chunks=int(payload["total_chunks"]),
            index_path=str(payload["index_path"]),
        )

    @app.get("/api/document/full_text")
    def get_full_text(source_path: str, record_index: int = 0) -> dict:
        try:
            import json
            from pathlib import Path
            path = Path(source_path)
            if not path.exists():
                return JSONResponse({"error": f"File not found: {source_path}"}, status_code=404)
            with open(path, "r", encoding="utf-8") as f:
                payload = json.load(f)
            records = payload if isinstance(payload, list) else [payload]
            if record_index < 0 or record_index >= len(records):
                return JSONResponse({"error": f"Record index {record_index} out of bounds"}, status_code=400)
            record = records[record_index]
            content = record.get("cleaned_content") or record.get("content") or ""
            return {"full_text": content}
        except Exception as e:
            return JSONResponse({"error": str(e)}, status_code=500)

    @app.get("/api/corpus/stats")
    def corpus_stats() -> dict:
        # Return simple corpus statistics and available sources/categories
        engine.ensure_index()
        total_docs = len({doc.extra.get("source_path") for doc in engine._documents})
        total_chunks = len(engine._documents)
        categories = {}
        sources = {}
        years = set()
        for doc in engine._documents:
            categories[doc.category] = categories.get(doc.category, 0) + 1
            sources[doc.source] = sources.get(doc.source, 0) + 1
            if isinstance(doc.year, int):
                years.add(doc.year)

        return {
            "total_documents": total_docs,
            "total_chunks": total_chunks,
            "categories": categories,
            "sources": sources,
            "year_range": [min(years) if years else None, max(years) if years else None],
        }

    def get_word_frequencies() -> dict[str, int]:
        import logging
        logger = logging.getLogger(__name__)
        freq_path = config.output_dir / "word_frequencies.json"
        index_path = config.index_path

        # If index has been updated more recently than frequencies cache, clear cache
        if freq_path.exists() and index_path.exists():
            if index_path.stat().st_mtime > freq_path.stat().st_mtime:
                try:
                    logger.info("Index updated since last word frequency computation. Clearing cached frequencies...")
                    freq_path.unlink()
                except Exception:
                    pass

        if freq_path.exists():
            try:
                with freq_path.open("r", encoding="utf-8") as handle:
                    return json.load(handle)
            except Exception:
                pass

        # Fallback: compute from engine documents
        from collections import Counter
        engine.ensure_index()
        logger.info(f"Computing cached word frequencies from {len(engine._documents)} documents...")
        counter = Counter()
        for doc in engine._documents:
            words = [w.lower() for w in re.findall(r"[A-Za-z0-9]+", doc.text) if len(w) > 3]
            counter.update(words)
        for w in list(counter.keys()):
            if w in HIGHLIGHT_STOPWORDS:
                del counter[w]

        frequencies = {t: c for t, c in counter.most_common(500)}
        try:
            with freq_path.open("w", encoding="utf-8") as handle:
                json.dump(frequencies, handle, indent=2, ensure_ascii=True)
        except Exception:
            pass
        return frequencies

    def generate_wordcloud_bytes(top_n_local: int = 80) -> io.BytesIO | None:
        """Generate a PNG wordcloud from the current engine documents and return BytesIO."""
        try:
            from wordcloud import WordCloud
        except Exception:
            return None

        freq = get_word_frequencies()
        most = dict(sorted(freq.items(), key=lambda x: x[1], reverse=True)[:top_n_local])
        if not most:
            return None

        wc = WordCloud(width=1600, height=550, background_color='white', random_state=42)
        wc.generate_from_frequencies(most)

        img = wc.to_image()
        buf = io.BytesIO()
        img.save(buf, format='PNG')
        buf.seek(0)
        return buf

    @app.post("/api/generate_summary")
    def generate_summary(request: SummaryRequest) -> dict:
        from .energy_tracker import EnergyTracker
        with EnergyTracker("RAG Custom Summary Synthesis"):
            hits = engine.search(
                request.query,
                top_k=max(8, request.top_k),
                category=request.category,
                source=request.source,
                year=request.year,
            )
            if not hits:
                return {
                    "summary": "No relevant documents were found to synthesize a summary from.",
                    "hits": []
                }
            
            # Determine summary length guidelines
            if request.summary_length == "1":
                len_instruction = "Write a concise summary of approximately 300 words that fits on a single page."
            elif request.summary_length == "2":
                len_instruction = "Write a comprehensive and detailed report of approximately 750 words. To ensure the report is thorough enough to cover a minimum of 2 pages of a formal document, analyze the evidence deeply, detail specific operational rules, and expand on all key areas."
            else: # "3+"
                len_instruction = "Write an exhaustive, highly detailed policy briefing paper of approximately 1100 words. To ensure the report covers 3+ pages, provide a deep policy analysis, detail specific legal/regulatory implications, state-by-state variations, and implementation challenges."

            # Determine summary type guidelines
            if request.summary_type == "abstractive":
                type_instruction = (
                    "Type: Abstractive Summary.\n"
                    "Focus on synthesizing the retrieved evidence into a cohesive narrative. "
                    "Integrate all facts and findings smoothly. Focus on the core meaning and big picture."
                )
            elif request.summary_type == "comprehensive":
                type_instruction = (
                    "Type: Comprehensive Technical Report.\n"
                    "Focus on strict technical details, structures, and systems. "
                    "Use formal markdown headings: Overview, Detailed Regulatory Assessment, Conservation Impacts, and Implementation Framework. "
                    "Incorporate technical data, specific clauses, and numbers from the evidence."
                )
            elif request.summary_type == "evolution":
                type_instruction = (
                    "Type: Policy Evolution Analyst.\n"
                    "Focus on chronological development and baseline shifts over time. "
                    "Examine the document years/time periods of the evidence (e.g. from 2011 to 2026). "
                    "Highlight how earlier rules set the stage, what changes occurred in intermediate years, and what the latest documents show as the current operational baseline."
                )
            else: # "executive"
                type_instruction = (
                    "Type: Executive Briefing.\n"
                    "Focus on high-level strategic takeaways, action items, and confidence ratings. "
                    "Use a professional briefing tone, providing a concise summary first, followed by clear bulleted findings, critical challenges, and strategic recommendations."
                )

            # Build prompt
            prompt_lines = [
                "You are an expert wildlife policy analyst and legal researcher.",
                "Your goal is to answer the user query based ONLY on the provided retrieved evidence snippets.",
                "Do not invent facts, and do not make assertions that cannot be traced to the evidence.",
                "Structure your output cleanly with markdown headings matching the selected style.",
                "Always cite sources inline using square brackets with the document titles (e.g., [National Wildlife Action Plan 2017-2031]).",
                "",
                f"LENGTH INSTRUCTION: {len_instruction}",
                f"STYLE INSTRUCTION:\n{type_instruction}",
                "",
                "Retrieved Evidence Snippets:",
            ]
            for idx, hit in enumerate(hits[:8], start=1):
                snippet = " ".join(hit.get('text', '').split())
                title = hit.get('title') or f"doc{idx}"
                year = hit.get('year') or 'N/A'
                src = hit.get('source') or ''
                prompt_lines.append(f"Snippet [{idx}] (Title: {title} ({year}) | Source: {src}):\n{snippet}\n")
            
            prompt_lines.extend([
                "",
                "User query:",
                request.query,
                "",
                "Write the detailed report now:"
            ])
            prompt = "\n".join(prompt_lines)

            try:
                summary_text = ollama_generate(prompt, model="llama3.2:3b", timeout=90)
            except Exception as exc:
                summary_text = engine.answer(request.query, hits)

            return {
                "summary": summary_text,
                "hits": hits
            }

    def find_matched_animal_image(query: str) -> str | None:
        q = query.lower()
        species_dir = Path("data/dataset/images/species")
        if not species_dir.exists():
            return None
        
        files = list(species_dir.glob("*.jpg"))
        for f in files:
            name_clean = f.stem.replace("-", " ")
            if name_clean in q:
                return str(f)
        
        if "tiger" in q:
            return str(species_dir / "bengal-tiger.jpg")
        if "elephant" in q:
            return str(species_dir / "indian-elephant.jpg")
        if "leopard" in q:
            if "snow" in q:
                return str(species_dir / "snow-leopard.jpg")
            return str(species_dir / "clouded-leopard.jpg")
        if "cheetah" in q:
            return str(species_dir / "asiatic-cheetah.jpg")
        if "dolphin" in q:
            return str(species_dir / "ganges-river-dolphin.jpg")
        if "rhino" in q or "rhinoceros" in q:
            return str(species_dir / "greater-one-horned-rhinoceros.jpg")
        if "bear" in q:
            return str(species_dir / "sloth-bear.jpg")
        if "eagle" in q:
            return str(species_dir / "grey-headed-fish-eagle.jpg")
        if "monkey" in q or "macaque" in q:
            return str(species_dir / "rhesus-macaque.jpg")
            
        animal_terms = {"animal", "animals", "species", "wildlife", "fauna", "biodiversity", "zoo", "zoos"}
        if any(t in q for t in animal_terms):
            return str(species_dir / "bengal-tiger.jpg")
            
        return None

    def generate_visualization_chart(hits) -> str | None:
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            import tempfile
            
            if not hits:
                return None
                
            titles = [h.get('title', '')[:30] + "..." if len(h.get('title', '')) > 30 else h.get('title', '') for h in hits[:5]]
            raw_scores = [float(h.get('score', 0)) for h in hits[:5]]
            
            scores = []
            for s in raw_scores:
                if s < 0:
                    scores.append(abs(s))
                else:
                    scores.append(s)
            
            fig, ax = plt.subplots(figsize=(6, 2.8))
            colors = ['#132a24', '#1f4e43', '#2a7262', '#369781', '#41bc9f']
            bars = ax.barh(titles[::-1], scores[::-1], color=colors[:len(hits)][::-1], edgecolor='#132a24', height=0.55)
            
            ax.set_title("Document Retrieval Relevance Details", fontsize=10, fontweight='bold', color='#132a24', pad=10)
            ax.set_xlabel("Relevance Score (higher is more relevant)", fontsize=8, color='#132a24')
            ax.tick_params(axis='both', which='major', labelsize=8)
            
            for spine in ('top', 'right'):
                ax.spines[spine].set_visible(False)
            ax.spines['left'].set_color('#132a24')
            ax.spines['bottom'].set_color('#132a24')
            
            plt.tight_layout()
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            plt.savefig(tmp.name, dpi=150, bbox_inches='tight')
            plt.close()
            return tmp.name
        except Exception as e:
            logger.error(f"Error generating matplotlib chart: {e}")
            return None

    from reportlab.pdfgen import canvas
    from reportlab.lib.colors import HexColor

    class WildaiCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, total_pages):
            self.saveState()
            width, height = self._pagesize
            
            # 1. Double Line Border
            self.setStrokeColor(HexColor("#132a24"))
            self.setLineWidth(1)
            self.rect(36, 36, width - 72, height - 72)
            self.rect(39, 39, width - 78, height - 78)
            
            # 2. Diagonal Watermark "WILDAI"
            self.setFont("Helvetica-Bold", 70)
            self.setFillColor(HexColor("#132a24"), alpha=0.035)
            self.saveState()
            self.translate(width / 2, height / 2)
            self.rotate(45)
            self.drawCentredString(0, 0, "WILDAI")
            self.restoreState()
            
            # 3. Header
            self.setFont("Helvetica-Bold", 8)
            self.setFillColor(HexColor("#132a24"))
            self.drawString(50, height - 28, "WILDAI RESEARCH & PLANNING CONSOLE")
            self.setFont("Helvetica", 8)
            self.setFillColor(HexColor("#64748b"))
            self.drawRightString(width - 50, height - 28, "Wildlife Intelligence RAG System")
            
            self.setStrokeColor(HexColor("#cbd5e1"))
            self.setLineWidth(0.5)
            self.line(40, height - 32, width - 40, height - 32)
            
            # 4. Footer
            page_num = self._pageNumber
            self.drawRightString(width - 50, 24, f"Page {page_num} of {total_pages}")
            self.drawString(50, 24, "Confidential · Document Citations & Passage Reference Logs")
            self.line(40, 30, width - 40, 30)
            
            self.restoreState()

    @app.post("/api/export")
    def export_result(request: ExportRequest, fmt: str = "md") -> Any:
        from .energy_tracker import EnergyTracker
        with EnergyTracker(f"Query Result Exporting ({fmt.upper()})"):
            hits = engine.search(request.query, top_k=request.top_k, category=request.category, source=request.source, year=request.year)
            answer = engine.answer(request.query, hits)

            timestamp = datetime.datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
            filename = f"wildai_report_{timestamp}.{fmt}"

            if fmt == "md":
                body_lines = [
                    "# WILDAI Query Report", 
                    "", 
                    f"**Query:** {request.query}", 
                    "", 
                    "## I. Executive & Analytical Summary",
                    "",
                    request.detailed_report or answer,
                    "",
                    "## II. Grounded Reference Passages & Sources",
                    ""
                ]
                
                if request.attach_snippets:
                    for idx, hit in enumerate(hits[:5], start=1):
                        url_part = f"([link]({hit['url']}))" if hit.get('url') else ""
                        body_lines.append(f"### [{idx}] {hit['title']} ({hit.get('year')}) — {hit.get('source')} {url_part}")
                        body_lines.append(f"> {hit['text'][:800]}")
                        body_lines.append("")
                        
                content = "\n".join(body_lines)
                
                if request.include_animal_photo:
                    img_path = find_matched_animal_image(request.query)
                    if img_path:
                        try:
                            import base64
                            with open(img_path, "rb") as image_file:
                                encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
                            content = f"![species](data:image/jpeg;base64,{encoded_string})\n\n" + content
                        except Exception:
                            pass
                
                tmp = NamedTemporaryFile(delete=False, suffix=".md")
                tmp.write(content.encode("utf-8"))
                tmp.flush()
                return FileResponse(tmp.name, filename=filename, media_type="text/markdown")

            if fmt in {"pdf", "docx"}:
                if fmt == "pdf":
                    from reportlab.lib.pagesizes import letter
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

                    tmp = NamedTemporaryFile(delete=False, suffix=".pdf")
                    doc = SimpleDocTemplate(
                        tmp.name,
                        pagesize=letter,
                        leftMargin=54,
                        rightMargin=54,
                        topMargin=54,
                        bottomMargin=54
                    )
                    
                    styles = getSampleStyleSheet()
                    title_style = ParagraphStyle(
                        'ReportTitle',
                        parent=styles['Title'],
                        fontName='Helvetica-Bold',
                        fontSize=22,
                        leading=26,
                        textColor=HexColor('#132a24'),
                        alignment=TA_CENTER,
                        spaceAfter=15
                    )
                    
                    meta_style = ParagraphStyle(
                        'ReportMeta',
                        fontName='Helvetica-Oblique',
                        fontSize=9,
                        leading=12,
                        textColor=HexColor('#64748b'),
                        alignment=TA_CENTER,
                        spaceAfter=20
                    )
                    
                    h1_style = ParagraphStyle(
                        'Heading1Custom',
                        fontName='Helvetica-Bold',
                        fontSize=14,
                        leading=18,
                        textColor=HexColor('#132a24'),
                        spaceBefore=14,
                        spaceAfter=8,
                        keepWithNext=True
                    )

                    h2_style = ParagraphStyle(
                        'Heading2Custom',
                        fontName='Helvetica-Bold',
                        fontSize=12,
                        leading=16,
                        textColor=HexColor('#1f4e43'),
                        spaceBefore=10,
                        spaceAfter=6,
                        keepWithNext=True
                    )

                    body_style = ParagraphStyle(
                        'BodyTextCustom',
                        fontName='Helvetica',
                        fontSize=10,
                        leading=14.5,
                        textColor=HexColor('#1e293b'),
                        spaceAfter=10
                    )
                    
                    snippet_body_style = ParagraphStyle(
                        'SnippetBody',
                        fontName='Helvetica-Oblique',
                        fontSize=8.5,
                        leading=12,
                        textColor=HexColor('#334155'),
                    )

                    story = []
                    story.append(Paragraph("WILDAI Query Synthesis & Policy Report", title_style))
                    
                    timestamp_str = datetime.datetime.now().strftime("%d %B %Y, %I:%M %p")
                    story.append(Paragraph(f"Query: \"{request.query}\"<br/>Generated on {timestamp_str} · Verified Environment", meta_style))
                    story.append(Spacer(1, 10))
                    
                    img_path = None
                    if request.include_animal_photo:
                        img_path = find_matched_animal_image(request.query)
                    
                    if img_path:
                        try:
                            img_flowable = Image(img_path, width=350, height=200)
                            story.append(img_flowable)
                            caption_text = f"Figure 1: Matched Species Illustration ({os.path.basename(img_path).replace('.jpg','').replace('-',' ').title()})"
                            story.append(Paragraph(caption_text, ParagraphStyle('Caption', fontName='Helvetica-Oblique', fontSize=8, leading=10, textColor=HexColor('#64748b'), alignment=TA_CENTER, spaceBefore=4, spaceAfter=15)))
                        except Exception as e:
                            logger.error(f"Failed to embed animal photo: {e}")
                    elif request.include_telemetry_charts:
                        chart_path = generate_visualization_chart(hits)
                        if chart_path:
                            try:
                                img_flowable = Image(chart_path, width=400, height=186)
                                story.append(img_flowable)
                                caption_text = "Figure 1: RAG Retrieval Relevance and Similarity Distribution"
                                story.append(Paragraph(caption_text, ParagraphStyle('Caption', fontName='Helvetica-Oblique', fontSize=8, leading=10, textColor=HexColor('#64748b'), alignment=TA_CENTER, spaceBefore=4, spaceAfter=15)))
                            except Exception as e:
                                logger.error(f"Failed to embed matplotlib chart: {e}")
                    
                    story.append(Spacer(1, 10))
                    story.append(Paragraph("I. Executive & Analytical Summary", h1_style))
                    
                    report_text = request.detailed_report or answer
                    for section in report_text.split('\n'):
                        section = section.strip()
                        if not section:
                            continue
                        if section.startswith("### "):
                            story.append(Paragraph(section.replace("### ", ""), h2_style))
                        elif section.startswith("## "):
                            story.append(Paragraph(section.replace("## ", ""), h1_style))
                        elif section.startswith("# "):
                            story.append(Paragraph(section.replace("# ", ""), title_style))
                        else:
                            clean_section = section
                            clean_section = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", clean_section)
                            story.append(Paragraph(clean_section, body_style))
                            
                    story.append(Spacer(1, 15))
                    
                    if request.attach_snippets and hits:
                        story.append(Paragraph("II. Grounded Reference Passages & Sources", h1_style))
                        for idx, hit in enumerate(hits[:5], start=1):
                            title_url = f"<a href='{hit['url']}' color='#1f4e43'><u>{hit['title']}</u></a>" if hit.get('url') else hit['title']
                            year_str = f" ({hit['year']})" if hit.get('year') else ""
                            relevance_pct = int(hit.get('semantic_score', hit['score']) * 100)
                            if relevance_pct < 0:
                                relevance_pct = int(hit.get('semantic_score', 0.75) * 100)
                            ref_header = f"[{idx}] {title_url}{year_str} — {hit.get('source', 'Unknown Source')} [Relevance: {relevance_pct}%]"
                            
                            cell_text = f"<b>{ref_header}</b><br/><br/><i>Excerpt:</i> \"{hit['text'][:800]}\""
                            cell_p = Paragraph(cell_text, snippet_body_style)
                            
                            snippet_table = Table([[cell_p]], colWidths=[480])
                            snippet_table.setStyle(TableStyle([
                                ('BACKGROUND', (0,0), (-1,-1), HexColor('#f8fafc')),
                                ('BOX', (0,0), (-1,-1), 0.5, HexColor('#e2e8f0')),
                                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                                ('TOPPADDING', (0,0), (-1,-1), 8),
                                ('BOTTOMPADDING', (0,0), (-1,-1), 8),
                                ('LEFTPADDING', (0,0), (-1,-1), 10),
                                ('RIGHTPADDING', (0,0), (-1,-1), 10),
                            ]))
                            
                            story.append(snippet_table)
                            story.append(Spacer(1, 10))
                    
                    doc.build(story, canvasmaker=WildaiCanvas)
                    return FileResponse(tmp.name, filename=filename, media_type="application/pdf")

                if fmt == "docx":
                    from docx import Document
                    from docx.shared import Inches, Pt, RGBColor
                    from docx.enum.text import WD_ALIGN_PARAGRAPH

                    doc = Document()
                    sections = doc.sections
                    for sec in sections:
                        sec.top_margin = Inches(1)
                        sec.bottom_margin = Inches(1)
                        sec.left_margin = Inches(1)
                        sec.right_margin = Inches(1)
                    
                    title = doc.add_paragraph()
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = title.add_run('WILDAI Query Synthesis & Policy Report')
                    run.font.name = 'Arial'
                    run.font.size = Pt(22)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(19, 42, 36)
                    
                    meta = doc.add_paragraph()
                    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_meta = meta.add_run(f'Query: "{request.query}"\nGenerated via WILDAI Research Console')
                    run_meta.font.name = 'Arial'
                    run_meta.font.size = Pt(9)
                    run_meta.font.italic = True
                    run_meta.font.color.rgb = RGBColor(100, 116, 139)
                    
                    img_path = None
                    if request.include_animal_photo:
                        img_path = find_matched_animal_image(request.query)
                    
                    if img_path:
                        try:
                            doc.add_paragraph().add_run().add_picture(img_path, width=Inches(4.5))
                            caption = doc.add_paragraph()
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap_run = caption.add_run(f'Figure 1: Matched Species Illustration ({os.path.basename(img_path).replace(".jpg","").replace("-"," ").title()})')
                            cap_run.font.name = 'Arial'
                            cap_run.font.size = Pt(8)
                            cap_run.font.italic = True
                        except Exception:
                            pass
                    
                    doc.add_heading('I. Executive & Analytical Summary', level=1)
                    report_text = request.detailed_report or answer
                    for line in report_text.split('\n'):
                        line = line.strip()
                        if not line:
                            continue
                        if line.startswith("### "):
                            p = doc.add_paragraph()
                            run = p.add_run(line.replace("### ", ""))
                            run.font.bold = True
                            run.font.size = Pt(12)
                            run.font.color.rgb = RGBColor(31, 78, 67)
                        elif line.startswith("## ") or line.startswith("# "):
                            doc.add_heading(line.replace("## ", "").replace("# ", ""), level=2)
                        else:
                            p = doc.add_paragraph(line)
                            p.paragraph_format.line_spacing = 1.15
                    
                    if request.attach_snippets and hits:
                        doc.add_heading('II. Grounded Reference Passages & Sources', level=1)
                        for idx, hit in enumerate(hits[:5], start=1):
                            relevance_pct = int(hit.get('semantic_score', hit['score']) * 100)
                            if relevance_pct < 0:
                                relevance_pct = int(hit.get('semantic_score', 0.75) * 100)
                            doc.add_paragraph(f"[{idx}] {hit['title']} ({hit.get('year')}) — {hit.get('source')} (Relevance: {relevance_pct}%)")
                            p_excerpt = doc.add_paragraph()
                            run_ex = p_excerpt.add_run(f'Excerpt: "{hit["text"][:800]}"')
                            run_ex.font.italic = True
                            p_excerpt.paragraph_format.left_indent = Inches(0.25)
                            
                    tmp = NamedTemporaryFile(delete=False, suffix=".docx")
                    doc.save(tmp.name)
                    return FileResponse(tmp.name, filename=filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            return JSONResponse({"error": "Unsupported format"}, status_code=400)

    @app.post("/api/chat")
    def chat_endpoint(request: QueryRequest, session_id: str | None = None) -> dict:
        # Very small retrieval + generation endpoint
        hits = engine.search(request.query, top_k=request.top_k, category=request.category, source=request.source, year=request.year)
        # For now, generation uses engine.answer (simple synthesis). If an LLM is wired, delegate to it.
        answer = engine.answer(request.query, hits)
        return {"session_id": session_id or f"s-{int(datetime.datetime.utcnow().timestamp())}", "query": request.query, "answer": answer, "hits": hits}

    @app.post('/api/chat/ollama')
    def chat_ollama(request: QueryRequest, model: str | None = None, session_id: str | None = None) -> dict:
        """RAG chat endpoint that uses the local Ollama LLM for final generation.

        It retrieves top hits from the index, constructs a prompt with the
        evidence snippets, and calls the local Ollama service via the
        `ollama_client.generate` adapter.
        """
        from .energy_tracker import EnergyTracker
        chosen_model = model or "llama3.2:3b"
        
        with EnergyTracker(f"Local LLM Chat (Ollama - {chosen_model})"):
            hits = engine.search(request.query, top_k=max(6, request.top_k), category=request.category, source=request.source, year=request.year)

            # build prompt
            prompt_lines = [
                "You are an expert assistant that answers user queries using provided evidence.",
                "Generate a detailed, comprehensive, and well-structured answer.",
                "Use clear bullet points or numbered lists to break down different aspects of the information.",
                "Cite the source titles in square brackets after assertions (e.g., [Himachal Pradesh Forest Rules (2023)]).",
                "Do not hallucinate facts that are not present in the evidence.",
                "Analyze the years and time periods of each retrieved document. You must reference these years/periods in your analysis to outline the timeframe of active policies (e.g. covering the period from 2011 to 2026), even if the text itself doesn't explicitly label them as 'latest'. If a query asks for 'latest' or 'recent' information, treat the most recently dated documents in the evidence as the latest policies and summarize their contents accordingly.",
                "",
                "Evidence:",
            ]
            for idx, hit in enumerate(hits[:6], start=1):
                snippet = " ".join(hit.get('text', '').split())[:400]
                title = hit.get('title') or f"doc{idx}"
                year = hit.get('year') or ''
                src = hit.get('source') or ''
                prompt_lines.append(f"{idx}. {title} ({year}) — {src}: {snippet}")

            prompt_lines.extend(["", "User question:", request.query, "", "Answer:" ] )
            prompt = "\n".join(prompt_lines)

            try:
                llm_out = ollama_generate(prompt, model=chosen_model, timeout=60)
            except Exception as exc:
                # fallback to engine.answer when Ollama not available
                answer = engine.answer(request.query, hits)
                return {"session_id": session_id or f"s-{int(datetime.datetime.utcnow().timestamp())}", "query": request.query, "answer": answer, "hits": hits, "warning": str(exc)}

            return {"session_id": session_id or f"s-{int(datetime.datetime.utcnow().timestamp())}", "query": request.query, "answer": llm_out, "hits": hits}

    @app.get("/api/analytics/terms")
    def analytics_terms(top_n: int = 50) -> dict:
        # Compute term frequencies across document chunks (naive).
        # If `query` is provided, restrict to matching hits.
        from collections import Counter

        query = None
        # allow query via query param ?q=...
        # FastAPI maps unknown params via request args; use fallback to global terms
        # For simplicity we read from request args in the URL
        try:
            from fastapi import Request

            @app.get('/_')
            def _noop():
                return {}
        except Exception:
            pass

        engine.ensure_index()
        counter = Counter()
        for doc in engine._documents:
            words = [w.lower() for w in re.findall(r"[A-Za-z0-9]+", doc.text) if len(w) > 3]
            counter.update(words)
        # filter stopwords
        stop = HIGHLIGHT_STOPWORDS
        for w in list(counter.keys()):
            if w in stop:
                del counter[w]

        most = counter.most_common(top_n)
        return {"top_terms": most}

    @app.get('/api/analytics/category_counts')
    def analytics_category_counts() -> dict:
        engine.ensure_index()
        counts: dict[str, int] = {}
        for doc in engine._documents:
            counts[doc.category] = counts.get(doc.category, 0) + 1
        return {"category_counts": counts}

    @app.get('/api/analytics/time_series')
    def analytics_time_series(category: str | None = None) -> dict:
        # Return counts per year, optionally filtered by category
        engine.ensure_index()
        counts: dict[int, int] = {}
        for doc in engine._documents:
            if category and doc.category != category:
                continue
            if isinstance(doc.year, int):
                counts[doc.year] = counts.get(doc.year, 0) + 1
        # return sorted list of (year, count)
        series = sorted(list(counts.items()))
        return {"time_series": series}

    @app.get('/api/analytics/wordcloud')
    def analytics_wordcloud(top_n: int = 80) -> dict:
        freq = get_word_frequencies()
        most = sorted(freq.items(), key=lambda x: x[1], reverse=True)[:top_n]
        return {"words": [{"term": t, "count": c} for t, c in most]}

    @app.get('/api/analytics/policy_time_series')
    def analytics_policy_time_series(source: str | None = None) -> dict:
        """Scan files in data/dataset/zoos-policy and return counts per year grouped by source/authority."""
        policy_dir = Path('data/dataset/zoos-policy')
        if not policy_dir.exists():
            return {"time_series": []}

        series: dict[int, dict[str, int]] = {}
        for path in policy_dir.glob('*.json'):
            try:
                with path.open('r', encoding='utf-8') as fh:
                    payload = json.load(fh)
            except Exception:
                continue
            year = payload.get('year')
            src = payload.get('source') or payload.get('url')
            if source and src and source.lower() not in src.lower():
                continue
            if not isinstance(year, int):
                # try to extract from content
                content = (payload.get('content') or '')
                ym = re.search(r"(19|20)\d{2}", content)
                year = int(ym.group(0)) if ym else None
            if year is None:
                continue
            series.setdefault(year, {})
            series[year][src] = series[year].get(src, 0) + 1

        # convert to sorted list
        result = []
        for y in sorted(series.keys()):
            result.append({"year": y, "counts": series[y]})
        return {"time_series": result}

    @app.get('/api/analytics/wordcloud_image')
    def analytics_wordcloud_image(q: str = None, top_n: int = 80) -> Any:
        # Build frequency dict and render PNG via python-wordcloud
        try:
            from wordcloud import WordCloud
        except Exception as e:
            return JSONResponse({"error": f"Missing dependency: {e}"}, status_code=500)

        if q:
            # Generate query-specific word cloud from search results
            hits = engine.search(q, top_k=20)
            from collections import Counter
            counter = Counter()
            for hit in hits:
                words = [w.lower() for w in re.findall(r"[A-Za-z0-9]+", hit["text"]) if len(w) > 3]
                counter.update(words)
            for w in list(counter.keys()):
                if w in HIGHLIGHT_STOPWORDS:
                    del counter[w]
            freq = {t: c for t, c in counter.most_common(top_n)}
        else:
            # Fall back to entire corpus frequencies
            freq = get_word_frequencies()

        most = dict(sorted(freq.items(), key=lambda x: x[1], reverse=True)[:top_n])
        if not most:
            return JSONResponse({"error": "No terms to render"}, status_code=400)

        wc = WordCloud(width=1600, height=550, background_color='white', random_state=42)
        wc.generate_from_frequencies(most)

        img = wc.to_image()
        buf = io.BytesIO()
        img.save(buf, format='PNG')
        buf.seek(0)
        return StreamingResponse(buf, media_type='image/png')

    @app.get('/api/analytics/energy')
    def get_energy_analytics() -> dict:
        import re
        from pathlib import Path
        from .energy_tracker import get_system_specs, LOG_FILE
        log_path = LOG_FILE
        try:
            specs = get_system_specs()
        except Exception:
            specs = {
                "os": "Windows",
                "cpu": "Intel(R) Core(TM) i7-11800H CPU @ 2.30GHz",
                "gpu": "NVIDIA GeForce RTX 3050 Laptop GPU",
                "ram": "16 GB"
            }
            
        logs = []
        if log_path.exists():
            pattern = re.compile(
                r"\[(?P<timestamp>.*?)\] Task: (?P<task>.*?) \| "
                r"Duration: (?P<duration>[\d.]+)s \| CPU Mean Util: (?P<cpu_util>[\d.]+)% \| "
                r"CPU Power: (?P<cpu_w>[\d.]+)W \| GPU Power: (?P<gpu_w>[\d.]+)W \| "
                r"Energy: (?P<energy>[\d.]+) Wh"
            )
            try:
                with log_path.open("r", encoding="utf-8") as f:
                    for line in f:
                        match = pattern.search(line)
                        if match:
                            d = match.groupdict()
                            logs.append({
                                "timestamp": d["timestamp"],
                                "task": d["task"],
                                "duration": float(d["duration"]),
                                "cpu_util": float(d["cpu_util"]),
                                "cpu_w": float(d["cpu_w"]),
                                "gpu_w": float(d["gpu_w"]),
                                "energy": float(d["energy"]),
                            })
            except Exception as e:
                return {"error": str(e), "logs": [], "system_specs": specs}
                
        return {"logs": logs, "system_specs": specs}

    @app.get('/api/analytics/year_category')
    def year_category_breakdown() -> dict:
        from collections import defaultdict
        engine.ensure_index()
        result_dict = defaultdict(lambda: defaultdict(int))
        for doc in engine._documents:
            if doc.year is not None:
                result_dict[int(doc.year)][doc.category] += 1
        return {"year_category": {yr: dict(cats) for yr, cats in result_dict.items()}}

    @app.post("/api/query", response_model=QueryResponse)

    def query(request: QueryRequest) -> QueryResponse:
        from .energy_tracker import EnergyTracker
        with EnergyTracker("RAG Search & Synthesis Query"):
            hits = engine.search(
                request.query,
                top_k=request.top_k,
                category=request.category,
                source=request.source,
                year=request.year,
            )
            answer = engine.answer(request.query, hits)
            highlight_terms = extract_highlight_terms(request.query)
            return QueryResponse(
                query=request.query,
                answer=answer,
                total_hits=len(hits),
                highlight_terms=highlight_terms,
                hits=[SearchHit(**hit) for hit in hits],
            )

    return app


app = create_app()
